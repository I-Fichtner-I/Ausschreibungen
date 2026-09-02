"""Stufe 2: Textextraktion aus Vergabeunterlagen."""

from __future__ import annotations

from pathlib import Path

import docx as docxlib
import pytest
from fpdf import FPDF
from openpyxl import Workbook

from tender_ai.extraction import extract_document, extractor_for
from tender_ai.extraction.base import DocumentExtractor
from tender_ai.models.document import ExtractionStatus


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    for line in (
        "Vergabeunterlagen Monitore",
        "Angebotsfrist: 15.10.2036",
        "Position 1: 500 Stueck 27 Zoll",
    ):
        pdf.cell(0, 10, line, new_x="LMARGIN", new_y="NEXT")
    path = tmp_path / "lv.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    document = docxlib.Document()
    document.add_paragraph("Vergabeunterlagen")
    document.add_paragraph("Angebotsfrist: 15.10.2036")
    table = document.add_table(rows=2, cols=3)
    for index, value in enumerate(["Pos", "Artikel", "Menge"]):
        table.cell(0, index).text = value
    for index, value in enumerate(["1", "Monitor 27 Zoll", "500"]):
        table.cell(1, index).text = value
    path = tmp_path / "unterlagen.docx"
    document.save(str(path))
    return path


@pytest.fixture
def xlsx_file(tmp_path: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Leistungsverzeichnis"
    sheet.append(["Pos", "Artikel", "Menge", "Einheit"])
    sheet.append([1, "Monitor 27 Zoll", 500, "Stueck"])
    sheet.append([2, "Docking-Station", 200, "Stueck"])
    path = tmp_path / "lv.xlsx"
    workbook.save(str(path))
    return path


def test_pdf_text_and_metadata(pdf_file: Path):
    result = extract_document(pdf_file, "application/pdf")
    assert result.status is ExtractionStatus.OK
    assert result.extractor == "pdf"
    assert result.page_count == 1
    assert "Angebotsfrist: 15.10.2036" in result.text
    assert result.checksum_sha256 and result.size_bytes
    assert result.page_text(1) is not None and result.page_text(99) is None


def test_docx_paragraphs_and_table(docx_file: Path):
    result = extract_document(docx_file)
    assert result.status is ExtractionStatus.OK
    assert "Angebotsfrist" in result.text
    assert len(result.tables) == 1
    table = result.tables[0]
    assert table.header == ["Pos", "Artikel", "Menge"]
    assert table.rows == [["1", "Monitor 27 Zoll", "500"]]
    assert table.row_count == 1 and table.column_count == 3


def test_xlsx_sheet_becomes_table_and_text(xlsx_file: Path):
    result = extract_document(xlsx_file)
    assert result.status is ExtractionStatus.OK
    table = result.tables[0]
    assert table.section == "Leistungsverzeichnis"
    assert table.header == ["Pos", "Artikel", "Menge", "Einheit"]
    assert len(table.rows) == 2
    assert "Docking-Station" in result.text


def test_html_strips_noise_and_reads_table(tmp_path: Path):
    path = tmp_path / "seite.html"
    path.write_text(
        "<html><head><title>Bekanntmachung</title><style>b{}</style></head>"
        "<body><script>egal()</script><p>Lieferung von Monitoren</p>"
        "<table><tr><th>Pos</th><th>Menge</th></tr><tr><td>1</td><td>500</td></tr></table>"
        "</body></html>",
        encoding="utf-8",
    )
    result = extract_document(path, "text/html")
    assert result.status is ExtractionStatus.OK
    assert "Lieferung von Monitoren" in result.text
    assert "egal()" not in result.text  # Skript-Inhalt ist kein Dokumenttext
    assert result.metadata["title"] == "Bekanntmachung"
    assert result.tables[0].header == ["Pos", "Menge"]


def test_csv_becomes_table_with_detected_delimiter(tmp_path: Path):
    path = tmp_path / "liste.csv"
    path.write_text("Pos;Artikel;Menge\n1;Monitor;500\n2;Maus;100\n", encoding="utf-8")
    result = extract_document(path)
    assert result.tables[0].header == ["Pos", "Artikel", "Menge"]
    assert result.tables[0].rows == [["1", "Monitor", "500"], ["2", "Maus", "100"]]


def test_unsupported_type_is_reported_not_raised(tmp_path: Path):
    path = tmp_path / "anhang.zip"
    path.write_bytes(b"PK\x03\x04egal")
    result = extract_document(path, "application/zip")
    assert result.status is ExtractionStatus.UNSUPPORTED
    assert result.error and "zip" in result.error
    assert extractor_for(path, "application/zip") is None


def test_broken_file_is_reported_as_failed(tmp_path: Path):
    path = tmp_path / "kaputt.pdf"
    path.write_bytes(b"das ist kein PDF")
    result = extract_document(path, "application/pdf")
    assert result.status is ExtractionStatus.FAILED
    assert result.error


def test_missing_file_is_reported(tmp_path: Path):
    result = extract_document(tmp_path / "gibtsnicht.pdf", "application/pdf")
    assert result.status is ExtractionStatus.FAILED
    assert "nicht lesbar" in (result.error or "")


def test_document_without_text_is_marked_empty(tmp_path: Path):
    """Ein gescanntes PDF liefert keinen Text - das muss sichtbar sein."""
    pdf = FPDF()
    pdf.add_page()
    path = tmp_path / "scan.pdf"
    pdf.output(str(path))
    result = extract_document(path, "application/pdf")
    assert result.status is ExtractionStatus.EMPTY
    assert result.ocr_used is False


def test_text_limit_truncates_and_marks_partial(tmp_path: Path):
    path = tmp_path / "lang.txt"
    path.write_text("A" * 5000, encoding="utf-8")
    extractor = extractor_for(path, "text/plain", max_characters=100)
    assert isinstance(extractor, DocumentExtractor)
    result = extractor.extract(path, "text/plain")
    assert result.truncated is True
    assert result.status is ExtractionStatus.PARTIAL
    assert result.character_count == 100


def test_media_type_wins_over_suffix(tmp_path: Path):
    """Der gemeldete Typ entscheidet - die Endung ist nur Rueckfallebene."""
    path = tmp_path / "daten.bin"
    path.write_text("Nur Text", encoding="utf-8")
    assert extract_document(path, "text/plain").extractor == "text"
    assert extract_document(path).status is ExtractionStatus.UNSUPPORTED


def test_pdf_table_is_detected_with_header(tmp_path: Path):
    """Tabellen im PDF sind das Rohmaterial der Artikelerkennung (Stufe 3)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=10)
    with pdf.table() as table:
        for row in (
            ["Pos", "Artikel", "Menge"],
            ["1", "Monitor 27 Zoll", "500"],
            ["2", "Docking-Station", "200"],
        ):
            table_row = table.row()
            for cell in row:
                table_row.cell(cell)
    path = tmp_path / "mit_tabelle.pdf"
    pdf.output(str(path))

    result = extract_document(path, "application/pdf")
    assert len(result.tables) == 1
    detected = result.tables[0]
    assert detected.page == 1
    assert detected.header == ["Pos", "Artikel", "Menge"]
    assert detected.rows == [["1", "Monitor 27 Zoll", "500"], ["2", "Docking-Station", "200"]]


@pytest.mark.parametrize(
    ("rows", "expect_header"),
    [
        ([["Pos", "Artikel"], ["1", "Monitor"]], True),
        ([["1", "2"], ["3", "4"]], False),  # nur Zahlen -> Datenzeile, keine Kopfzeile
        ([["Pos", ""], ["1", "Monitor"]], False),  # Luecke -> unsicher, lieber keine
        ([["Pos", "Artikel"]], False),  # einzelne Zeile ist kein Kopf
    ],
)
def test_header_detection_is_conservative(rows, expect_header):
    """Lieber keine Kopfzeile annehmen, als eine Datenzeile zu verlieren."""
    from tender_ai.extraction.pdf import _to_table

    table = _to_table([list(row) for row in rows], page=1, index=0)
    assert table is not None
    assert (table.header is not None) is expect_header
    expected_rows = len(rows) - 1 if expect_header else len(rows)
    assert table.row_count == expected_rows


def test_empty_pdf_table_is_dropped():
    from tender_ai.extraction.pdf import _to_table

    assert _to_table([["", None], [" ", ""]], page=1, index=0) is None
