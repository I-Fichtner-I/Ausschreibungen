"""DOCX-Extraktion mit python-docx (Absaetze und Tabellen)."""

from __future__ import annotations

from pathlib import Path
from typing import ClassVar

import docx

from ..models.document import ExtractedDocument, ExtractedPage, ExtractedTable
from .base import DocumentExtractor, register_extractor

_DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@register_extractor
class DocxExtractor(DocumentExtractor):
    name: ClassVar[str] = "docx"
    media_types: ClassVar[tuple[str, ...]] = (_DOCX_MEDIA, "application/msword")
    suffixes: ClassVar[tuple[str, ...]] = (".docx",)

    def _extract(self, path: Path, document: ExtractedDocument) -> None:
        source = docx.Document(str(path))

        properties = source.core_properties
        document.metadata = {
            key: str(value)
            for key, value in (
                ("title", properties.title),
                ("author", properties.author),
                ("subject", properties.subject),
                ("created", properties.created),
                ("modified", properties.modified),
            )
            if value
        }

        # DOCX kennt keine Seiten (die entstehen erst beim Rendern) - der Text
        # wird deshalb als eine Seite gefuehrt.
        paragraphs = [p.text.strip() for p in source.paragraphs if p.text.strip()]
        document.pages.append(ExtractedPage(number=1, text="\n".join(paragraphs)))

        for index, table in enumerate(source.tables):
            rows = [
                [cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows
            ]
            rows = [row for row in rows if any(cell for cell in row)]
            if not rows:
                continue
            header = rows[0] if len(rows) > 1 and all(cell for cell in rows[0]) else None
            document.tables.append(
                ExtractedTable(
                    page=1, index=index, header=header, rows=rows[1:] if header else rows
                )
            )
